"""Tests for agent.memory.doc_reader — document text extraction."""

import os
import tempfile

import pytest

from agent.memory.doc_reader import extract_text, SUPPORTED_EXTENSIONS


# ── TXT extraction ──────────────────────────────────────────────────────────


def test_extract_txt():
    """Plain text file is read as-is."""
    content = "Hello, world!\nThis is a test."
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False,
                                     encoding="utf-8") as f:
        f.write(content)
        path = f.name
    try:
        result = extract_text(path)
        assert result is not None
        assert result["text"] == content
        assert result["word_count"] > 0
        assert result["extension"] == ".txt"
    finally:
        os.unlink(path)


def test_extract_txt_unicode():
    """Unicode text is preserved correctly."""
    content = "Hëllö Wörld ☃️ — café résumé"
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False,
                                     encoding="utf-8") as f:
        f.write(content)
        path = f.name
    try:
        result = extract_text(path)
        assert result is not None
        assert "Hëllö" in result["text"]
        assert "☃️" in result["text"]
    finally:
        os.unlink(path)


def test_extract_txt_empty():
    """Empty text file returns empty text with 0 word count."""
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False,
                                     encoding="utf-8") as f:
        path = f.name
    try:
        result = extract_text(path)
        assert result is not None
        assert result["text"] == "" or result["text"].strip() == ""
        assert result["word_count"] == 0
    finally:
        os.unlink(path)


# ── MD extraction ───────────────────────────────────────────────────────────


def test_extract_md():
    """Markdown file returns content with text, stripped of formatting."""
    md_content = """# Title

This is a **bold** paragraph with a [link](http://example.com).

- List item 1
- List item 2

```python
print("hello")
```
"""
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False,
                                     encoding="utf-8") as f:
        f.write(md_content)
        path = f.name
    try:
        result = extract_text(path)
        assert result is not None
        # Core text content should be present
        assert "Title" in result["text"]
        assert "bold" in result["text"]
        assert "List item 1" in result["text"]
        assert result["extension"] == ".md"
    finally:
        os.unlink(path)


# ── PDF extraction ──────────────────────────────────────────────────────────


def test_extract_pdf_simple():
    """Simple PDF text is extracted correctly."""
    pytest.importorskip("pypdf")
    canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name
    try:
        c = canvas.Canvas(path)
        c.drawString(100, 750, "Hello PDF World")
        c.drawString(100, 730, "Page 1 content here.")
        c.save()

        result = extract_text(path)
        assert result is not None
        assert "Hello PDF World" in result["text"]
        assert result["extension"] == ".pdf"
        assert result["word_count"] >= 4
    finally:
        os.unlink(path)


def test_extract_pdf_multi_page():
    """Multi-page PDF extracts text from all pages."""
    pytest.importorskip("pypdf")
    canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name
    try:
        c = canvas.Canvas(path)
        c.drawString(100, 750, "Page One Content")
        c.showPage()
        c.drawString(100, 750, "Page Two Content")
        c.save()

        result = extract_text(path)
        assert result is not None
        assert "Page One Content" in result["text"]
        assert "Page Two Content" in result["text"]
    finally:
        os.unlink(path)


def test_extract_pdf_empty():
    """Empty PDF (no text) returns empty result, not error."""
    pytest.importorskip("pypdf")
    canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name
    try:
        c = canvas.Canvas(path)
        c.save()

        result = extract_text(path)
        # Should not raise — returns whatever text is found
        assert result is not None
        assert result["extension"] == ".pdf"
    finally:
        os.unlink(path)


# ── DOCX extraction ─────────────────────────────────────────────────────────


def test_extract_docx_simple():
    """Simple DOCX text is extracted correctly."""
    pytest.importorskip("docx")
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        doc = Document()
        doc.add_paragraph("Hello DOCX World")
        doc.add_paragraph("Second paragraph here.")
        doc.save(path)

        result = extract_text(path)
        assert result is not None
        assert "Hello DOCX World" in result["text"]
        assert "Second paragraph" in result["text"]
        assert result["extension"] == ".docx"
    finally:
        os.unlink(path)


def test_extract_docx_with_formatting():
    """DOCX with bold/italic/headings extracts plain text content."""
    pytest.importorskip("docx")
    from docx import Document
    from docx.shared import Pt
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        doc = Document()
        doc.add_heading("Title Heading", level=1)
        p = doc.add_paragraph()
        run = p.add_run("Bold text")
        run.bold = True
        p.add_run(" and normal text")
        doc.add_paragraph("Final line.")
        doc.save(path)

        result = extract_text(path)
        assert result is not None
        assert "Title Heading" in result["text"]
        assert "Bold text" in result["text"]
        assert "normal text" in result["text"]
        assert "Final line" in result["text"]
    finally:
        os.unlink(path)


# ── Error handling ──────────────────────────────────────────────────────────


def test_extract_unsupported_extension():
    """Unsupported extension returns None."""
    with tempfile.NamedTemporaryFile(suffix=".xyz", mode="w", delete=False) as f:
        f.write("some data")
        path = f.name
    try:
        result = extract_text(path)
        assert result is None
    finally:
        os.unlink(path)


def test_extract_nonexistent_file():
    """Non-existent file returns None without raising."""
    result = extract_text("/nonexistent/path/file.txt")
    assert result is None


def test_extract_binary_file_with_txt_ext():
    """Binary content with .txt extension does not crash."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
        f.write(b"\x00\x01\x02\xff\xfe")
        path = f.name
    try:
        result = extract_text(path)
        # Should handle gracefully — may contain replacement chars
        assert result is not None
        assert "text" in result
    finally:
        os.unlink(path)


# ── SUPPORTED_EXTENSIONS ────────────────────────────────────────────────────


def test_supported_extensions_contains_expected():
    """SUPPORTED_EXTENSIONS includes all handled document types."""
    assert ".txt" in SUPPORTED_EXTENSIONS
    assert ".md" in SUPPORTED_EXTENSIONS
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS


def test_extract_txt_very_large_truncation():
    """Very large text file is truncated to MAX_CHAR_LIMIT."""
    from agent.memory.doc_reader import MAX_CHAR_LIMIT
    large_content = "word " * (MAX_CHAR_LIMIT + 10000)
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False,
                                     encoding="utf-8") as f:
        f.write(large_content)
        path = f.name
    try:
        result = extract_text(path)
        assert result is not None
        assert len(result["text"]) <= MAX_CHAR_LIMIT
    finally:
        os.unlink(path)
